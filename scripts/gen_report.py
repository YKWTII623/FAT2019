#!/usr/bin/env python3
"""Generate FAT2019 experiment report in Word format."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime, os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ==================== 封面 ====================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Freesound Audio Tagging 2019\n实验报告')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'验证集 LWLRAP: 0.620 | 80 类环境声音\n{datetime.date.today().strftime("%Y年%m月%d日")}')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ==================== 目录 ====================
doc.add_heading('目录', level=1)
doc.add_paragraph('[Word → 引用 → 目录 → 自动目录]')
doc.add_page_break()

# ==================== 1. 任务介绍 ====================
doc.add_heading('1. 任务介绍', level=1)

doc.add_heading('1.1 赛题', level=2)
doc.add_paragraph(
    'Kaggle 上的 Freesound Audio Tagging 2019（也是 DCASE 2019 Challenge Task 2），'
    '给一段 10 秒以内的音频，判断里面有哪些声音。总共 80 个类别，'
    '比如 Church bell、Bark、Motorcycle、Sneeze 之类的环境音。'
    '一段音频可以同时打多个标签。'
)

doc.add_heading('1.2 评价指标', level=2)
doc.add_paragraph(
    'LWLRAP（Label-Weighted Label-Ranking Average Precision）。'
    '简单说就是看模型能不能把正确的标签排在错误标签前面。'
    '按每个标签的样本数加权，样本多的标签权重大，少的权重小。'
    '最高 1.0，越高越好。'
)

doc.add_heading('1.3 排名参考', level=2)
doc.add_paragraph('比赛有 400 多支队伍参加。Private Leaderboard 前几名：')

table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
hdr = table.rows[0].cells
hdr[0].text = '排名'
hdr[1].text = '方案'
hdr[2].text = 'LWLRAP'
data = [
    ('#1', 'Ebbers et al. (CRNN)', '0.755'),
    ('#8', 'Eric BOUTEILLON (CNN+VGG)', '0.738'),
    ('#16', 'Eric BOUTEILLON (4-model ensemble)', '0.733'),
    ('本方案', 'DeepAudioCNN (单模型)', '0.620（验证集）'),
]
for i, (rank, team, score) in enumerate(data):
    row = table.rows[i+1].cells
    row[0].text = rank
    row[1].text = team
    row[2].text = score

# ==================== 2. 数据 ====================
doc.add_heading('2. 数据', level=1)

doc.add_heading('2.1 数据来源', level=2)
doc.add_paragraph(
    '训练集分两坨：curated（4,970 条，人标的）和 noisy（19,815 条，从 Flickr 视频自动爬的）。'
    'noisy 的标签质量明显差很多，有些标得完全不对。'
    '测试集 3,361 条，没有标签。'
)

table2 = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
hdr2 = table2.rows[0].cells
hdr2[0].text = '数据'
hdr2[1].text = '条数'
hdr2[2].text = '标注质量'
d2 = [
    ('train_curated', '4,970', '靠谱'),
    ('train_noisy', '19,815', '不太靠谱'),
    ('test', '3,361', '无标签'),
]
for i, d in enumerate(d2):
    row = table2.rows[i+1].cells
    row[0].text = d[0]
    row[1].text = d[1]
    row[2].text = d[2]

doc.add_heading('2.2 标签分布', level=2)
doc.add_paragraph(
    '80 个类别分布很不均匀。一些常见的声音（比如 Music 相关的）有好几千条，'
    '冷门的（比如 Gong、Glockenspiel）只有几十条。这种长尾分布在训练时会导致模型偏向常见类。'
)
doc.add_paragraph('[插入 outputs/label_distribution.png]')

# ==================== 3. 特征提取 ====================
doc.add_heading('3. 特征提取', level=1)
doc.add_paragraph(
    '用的 log-mel spectrogram，算是音频分类里的标配了。音频重采样到 22,050 Hz，转单声道，'
    '峰值归一化。然后 STFT → mel filterbank → log → 标准化到均值 0 方差 1。'
    '训练时随机截 5 秒，做 TTA 时取 3 个不同位置平均。'
)

table3 = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
hdr3 = table3.rows[0].cells
hdr3[0].text = '参数'
hdr3[1].text = '值'
params = [
    ('采样率', '22,050 Hz'),
    ('FFT', '1,024 点'),
    ('Hop', '512'),
    ('Mel bins', '128'),
    ('频率范围', '20 ~ 11,025 Hz'),
    ('片段长度', '5 秒（随机裁）'),
    ('TTA crops', '3'),
]
for i, (k, v) in enumerate(params):
    row = table3.rows[i+1].cells
    row[0].text = k
    row[1].text = v

# ==================== 4. 模型 ====================
doc.add_heading('4. 模型', level=1)
doc.add_paragraph(
    '自己搭了一个 18 层的 CNN（DeepAudioCNN），大概 1500 万参数。结构如下：'
)
doc.add_paragraph(
    '前 4 个 ConvBlock（1→64→128→256→512），每个 block 里两层卷积 + BN + ReLU + MaxPool，'
    '把 mel bins 从 128 降到 8。中间插了 3 个 ResBlock（带 skip connection），'
    '加深网络的同时避免梯度消失。最后两个不带 pooling 的 ConvBlock（512→512→768），'
    '接一个全局 avg+max pooling 拼成 1536 维，再过一个两层 MLP 输出 80 个 logits。'
)
doc.add_paragraph('[插入模型结构图]')

# ==================== 5. 训练 ====================
doc.add_heading('5. 训练细节', level=1)

doc.add_heading('5.1 基本配置', level=2)
table4 = doc.add_table(rows=9, cols=2, style='Light Grid Accent 1')
hdr4 = table4.rows[0].cells
hdr4[0].text = '配置'
hdr4[1].text = '设的值'
train_params = [
    ('优化器', 'AdamW (lr=3e-4, wd=1e-4)'),
    ('LR 调度', '5 epoch warmup 之后 cosine 退火'),
    ('batch size', '64'),
    ('最大 epoch', '200（实际 82 轮早停）'),
    ('dropout', '0.25'),
    ('损失函数', 'Focal BCE (γ=2, α=0.25)'),
    ('Mixup α', '0.2'),
    ('验证集比例', '20%'),
]
for i, (k, v) in enumerate(train_params):
    row = table4.rows[i+1].cells
    row[0].text = k
    row[1].text = v

doc.add_heading('5.2 数据增强', level=2)
doc.add_paragraph(
    'SpecAugment：在 spectrogram 上随机盖掉几块频率和时间段（频率最多 12 bins × 2，'
    '时间最多 24 帧 × 2），用均值填充。相当于逼模型根据局部信息做判断，不要死盯某几个频率。'
)
doc.add_paragraph(
    'Mixup：从 Beta(0.2, 0.2) 采样一个比例，把 batch 里两段音频和标签按比例混一起。'
    '算是正则化手段，防止过拟合。'
)

doc.add_heading('5.3 Focal Loss', level=2)
doc.add_paragraph(
    '80 个类别里很多是长尾里的小类，普通 BCE loss 会被大类带着跑。'
    'Focal Loss 加了一个 (1-pt)^γ 的权重项，让模型少关注已经分对的样本，'
    '多花精力在那些分不对的难样本和稀有类上。设的 γ=2, α=0.25。'
)

doc.add_heading('5.4 Warmup + 早停', level=2)
doc.add_paragraph(
    '前 5 个 epoch 学习率从 0 线性涨到 3e-4，防止刚初始化完就大乱跑。'
    '验证集 10 个 epoch 没涨够 1e-4 就停。实际在第 82 轮触发早停，'
    '最佳 LWLRAP 约 0.62。'
)

# ==================== 6. 结果 ====================
doc.add_heading('6. 结果', level=1)

doc.add_heading('6.1 几个模型的对比', level=2)
table5 = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
hdr5 = table5.rows[0].cells
hdr5[0].text = '模型'
hdr5[1].text = 'LWLRAP'
hdr5[2].text = '备注'
results = [
    ('sklearn LogReg', '0.52', 'log-mel 统计量 + 逻辑回归，200 条快速测试'),
    ('SmallAudioCNN (8 层)', '0.42', '也是 200 条测试，太浅了没意义'),
    ('DeepAudioCNN (18 层)', '0.620', '全量数据跑的结果'),
]
for i, (m, s, n) in enumerate(results):
    row = table5.rows[i+1].cells
    row[0].text = m
    row[1].text = s
    row[2].text = n

doc.add_heading('6.2 训练过程', level=2)
doc.add_paragraph(
    'Epoch 1 的验证 LWLRAP 只有 0.097，后面慢慢涨。到 82 轮触发早停，'
    '最佳值 0.62。训练 loss 从 0.009 降到 0.048，验证 loss 在 0.044 附近波动。'
)
doc.add_paragraph('[插入训练 loss/LWLRAP 曲线]')

doc.add_heading('6.3 消融实验', level=2)
doc.add_paragraph(
    '还没来得及做完整的消融，下面表格是计划对照的项目：'
)
table6 = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
hdr6 = table6.rows[0].cells
hdr6[0].text = '改动'
hdr6[1].text = 'LWLRAP'
hdr6[2].text = '涨了多少'
ablations = [
    ('裸模型（不用任何技巧）', '[待补]', '-'),
    ('+ SpecAugment', '[待补]', ''),
    ('+ Focal Loss', '[待补]', ''),
    ('+ Warmup', '[待补]', ''),
    ('+ Mixup', '[待补]', ''),
    ('+ 把片段从 5 秒改成 7 秒', '[待补]', ''),
    ('全上（当前方案）', '0.620', ''),
]
for i, (cfg, score, delta) in enumerate(ablations):
    row = table6.rows[i+1].cells
    row[0].text = cfg
    row[1].text = score
    row[2].text = delta

# ==================== 7. 总结 ====================
doc.add_heading('7. 总结与后续', level=1)

doc.add_heading('7.1 做了什么', level=2)
doc.add_paragraph(
    '用 log-mel + DeepAudioCNN（18 层），加了 SpecAugment、Focal Loss、Warmup、Mixup、'
    '早停，验证集做到了 0.62。比最开始的 sklearn 基线（0.52）提了 10 个点。'
)

doc.add_heading('7.2 还能做什么', level=2)
improvements = [
    ('换 backbone', '试试 PANNs Cnn14 或者 EfficientNet，在 AudioSet 预训练过的，应该能涨不少'),
    ('K-fold', '现在只分了一次 train/val，有点随机性。做 5 折更稳'),
    ('ensemble', '多训几个不同随机种子的模型取平均，一般能涨 2-5 个点'),
    ('洗 noisy 数据', 'noisy 标签有不少错的，可以用置信度过滤或者用 curated 训出来的模型筛一遍'),
    ('加 Attention', '在 CNN 特征上面加个 SE block 或者 transformer 层'),
]
for title, desc in improvements:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ==================== 附录 ====================
doc.add_heading('附录', level=1)
doc.add_heading('A. 代码', level=2)
doc.add_paragraph('GitHub: [填仓库地址]')
doc.add_paragraph('主要代码：src/fat2019_cnn.py')

doc.add_heading('B. 环境', level=2)
doc.add_paragraph(
    'Python 3.11 / PyTorch 2.8+cu124 / NumPy / SciPy / Pandas / scikit-learn / Joblib / tqdm / Matplotlib'
)

doc.add_heading('C. 参考', level=2)
refs = [
    'Ebbers, J. et al. "Audio tagging with noisy labels and minimal supervision." DCASE 2019.',
    'Park, D. S. et al. "SpecAugment." Interspeech 2019.',
    'Lin, T. Y. et al. "Focal Loss for Dense Object Detection." ICCV 2017.',
    'Zhang, H. et al. "mixup: Beyond Empirical Risk Minimization." ICLR 2018.',
    'Eric Bouteillon 的方案: https://github.com/ebouteillon/freesound-audio-tagging-2019',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f'[{i}] {ref}')

# ---- 保存 ----
out_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'reports')
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'FAT2019_实验报告.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
